# converters/word_converter.py
"""Word 转 Markdown 转换器"""
import mammoth
from docx import Document
from pathlib import Path
from typing import Optional, List
from converters.base_converter import BaseConverter
from utils.image_handler import ImageHandler

class WordConverter(BaseConverter):
    """Word 转 Markdown 转换器"""

    def convert(self, file_path: str) -> str:
        """
        将 Word 文件转换为 Markdown

        Args:
            file_path: Word 文件路径

        Returns:
            Markdown 内容
        """
        # 使用 python-docx 提取结构信息
        doc = Document(file_path)

        # 构建最终 Markdown
        markdown_parts = []

        # 处理段落，识别标题和正文
        for paragraph in doc.paragraphs:
            style_name = paragraph.style.name

            # 处理主标题（Title 样式或 Heading 0）
            if style_name == 'Title' or 'Heading 0' in style_name or style_name == '标题':
                markdown_parts.append(f"# {paragraph.text}\n")
            elif style_name.startswith('Heading') or style_name.startswith('标题'):
                # 处理标题
                level = self._get_heading_level(style_name)
                markdown_parts.append(f"\n{'#' * level} {paragraph.text}\n")
            else:
                # 处理普通段落
                if paragraph.text.strip():
                    markdown_parts.append(f"{paragraph.text}\n")

        # 处理表格
        tables_md = self._process_tables(doc.tables)
        if tables_md:
            markdown_parts.append("\n" + tables_md)

        return '\n'.join(markdown_parts)

    def _get_heading_level(self, style_name: str) -> int:
        """获取标题级别"""
        if 'Heading 1' in style_name or '标题 1' in style_name:
            return 2
        elif 'Heading 2' in style_name or '标题 2' in style_name:
            return 3
        elif 'Heading 3' in style_name or '标题 3' in style_name:
            return 4
        else:
            return 2  # 默认二级标题

    def _process_tables(self, tables: List) -> str:
        """处理表格"""
        if not tables:
            return ""

        markdown_parts = []

        for table in tables:
            rows = []

            # 提取表格数据
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(cells)

            if rows:
                # 生成 Markdown 表格
                # 表头
                header = rows[0]
                markdown_parts.append('| ' + ' | '.join(header) + ' |')
                markdown_parts.append('| ' + ' | '.join(['---'] * len(header)) + ' |')

                # 数据行
                for row in rows[1:]:
                    markdown_parts.append('| ' + ' | '.join(row) + ' |')

                markdown_parts.append('')

        return '\n'.join(markdown_parts)
