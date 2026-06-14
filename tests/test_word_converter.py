# tests/test_word_converter.py
"""Word 转换器测试"""
import pytest
import tempfile
from pathlib import Path
from docx import Document
from converters.word_converter import WordConverter

@pytest.fixture
def sample_word_file():
    """创建测试 Word 文件"""
    with tempfile.TemporaryDirectory() as tmpdir:
        doc = Document()
        doc.add_heading('测试文档', 0)
        doc.add_heading('第一章', level=1)
        doc.add_paragraph('这是一段测试文本。')
        doc.add_heading('第二章', level=1)
        doc.add_paragraph('这是另一段测试文本。')

        # 添加表格
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = '姓名'
        table.cell(0, 1).text = '年龄'
        table.cell(1, 0).text = '张三'
        table.cell(1, 1).text = '25'

        file_path = Path(tmpdir) / "test.docx"
        doc.save(str(file_path))

        yield str(file_path)

def test_convert_word_document(sample_word_file):
    """测试转换 Word 文档"""
    with tempfile.TemporaryDirectory() as tmpdir:
        converter = WordConverter(tmpdir)
        markdown = converter.convert(sample_word_file)

        # 检查是否包含标题
        assert "# 测试文档" in markdown
        assert "## 第一章" in markdown
        assert "## 第二章" in markdown

        # 检查是否包含段落
        assert "测试文本" in markdown

        # 检查是否包含表格
        assert "|" in markdown
        assert "张三" in markdown

def test_save_converted_word(sample_word_file):
    """测试保存转换后的 Word"""
    with tempfile.TemporaryDirectory() as tmpdir:
        converter = WordConverter(tmpdir)
        markdown = converter.convert(sample_word_file)

        # 保存文件
        output_path = converter.save_markdown(markdown, sample_word_file)

        # 验证文件存在
        assert output_path.exists()
        assert output_path.suffix == ".md"
